from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from services.local_ai_service import (
    generate_tailored_cv,
)


def add_list_section(
    document: Document,
    heading: str,
    values: list[str],
) -> None:
    """
    Add a heading followed by bullet-style paragraphs.
    """

    if not values:
        return

    document.add_heading(
        heading,
        level=1,
    )

    for value in values:
        cleaned_value = value.strip()

        if cleaned_value:
            document.add_paragraph(
                cleaned_value,
                style="List Bullet",
            )


def build_cv_text(
    result: dict,
) -> str:
    """
    Convert a structured CV result into plain text.
    """

    sections = []

    candidate_name = result.get(
        "candidate_name",
        "",
    ).strip()

    professional_title = result.get(
        "professional_title",
        "",
    ).strip()

    if candidate_name:
        sections.append(
            candidate_name
        )

    if professional_title:
        sections.append(
            professional_title
        )

    contact_details = result.get(
        "contact_details",
        [],
    )

    if contact_details:
        sections.append(
            " | ".join(contact_details)
        )

    summary = result.get(
        "professional_summary",
        "",
    ).strip()

    if summary:
        sections.append(
            f"PROFESSIONAL SUMMARY\n{summary}"
        )

    section_mapping = [
        (
            "TECHNICAL SKILLS",
            "technical_skills",
        ),
        (
            "PROFESSIONAL EXPERIENCE",
            "experience_sections",
        ),
        (
            "PROJECTS",
            "project_sections",
        ),
        (
            "EDUCATION",
            "education_sections",
        ),
        (
            "CERTIFICATIONS",
            "certification_sections",
        ),
        (
            "LANGUAGES",
            "language_sections",
        ),
        (
            "ADDITIONAL INFORMATION",
            "additional_sections",
        ),
    ]

    for heading, key in section_mapping:
        values = result.get(
            key,
            [],
        )

        formatted_values = "\n".join(
            f"- {value}"
            for value in values
            if value.strip()
        )

        if formatted_values:
            sections.append(
                f"{heading}\n{formatted_values}"
            )

    return "\n\n".join(sections)


def build_cv_docx(
    result: dict,
) -> bytes:
    """
    Convert the structured CV into a Word document.
    """

    document = Document()

    normal_style = document.styles[
        "Normal"
    ]

    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    candidate_name = result.get(
        "candidate_name",
        "",
    ).strip()

    professional_title = result.get(
        "professional_title",
        "",
    ).strip()

    if candidate_name:
        name_paragraph = document.add_paragraph()

        name_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        name_run = name_paragraph.add_run(
            candidate_name
        )

        name_run.bold = True
        name_run.font.size = Pt(18)

    if professional_title:
        title_paragraph = (
            document.add_paragraph()
        )

        title_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        title_run = title_paragraph.add_run(
            professional_title
        )

        title_run.bold = True
        title_run.font.size = Pt(12)

    contact_details = result.get(
        "contact_details",
        [],
    )

    if contact_details:
        contact_paragraph = (
            document.add_paragraph(
                " | ".join(contact_details)
            )
        )

        contact_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    summary = result.get(
        "professional_summary",
        "",
    ).strip()

    if summary:
        document.add_heading(
            "Professional Summary",
            level=1,
        )

        document.add_paragraph(
            summary
        )

    add_list_section(
        document=document,
        heading="Technical Skills",
        values=result.get(
            "technical_skills",
            [],
        ),
    )

    add_list_section(
        document=document,
        heading="Professional Experience",
        values=result.get(
            "experience_sections",
            [],
        ),
    )

    add_list_section(
        document=document,
        heading="Projects",
        values=result.get(
            "project_sections",
            [],
        ),
    )

    add_list_section(
        document=document,
        heading="Education",
        values=result.get(
            "education_sections",
            [],
        ),
    )

    add_list_section(
        document=document,
        heading="Certifications",
        values=result.get(
            "certification_sections",
            [],
        ),
    )

    add_list_section(
        document=document,
        heading="Languages",
        values=result.get(
            "language_sections",
            [],
        ),
    )

    add_list_section(
        document=document,
        heading="Additional Information",
        values=result.get(
            "additional_sections",
            [],
        ),
    )

    output = BytesIO()

    document.save(output)

    output.seek(0)

    return output.getvalue()


def create_tailored_cv(
    cv_text: str,
    job_text: str,
    job_details: dict,
    language: str,
) -> dict:
    """
    Generate a tailored CV and create TXT/DOCX versions.
    """

    result = generate_tailored_cv(
        cv_text=cv_text,
        job_text=job_text,
        job_details=job_details,
        language=language,
    )

    return {
        "structured_result": result,
        "text": build_cv_text(
            result
        ),
        "docx": build_cv_docx(
            result
        ),
        "warnings": result.get(
            "warnings",
            [],
        ),
    }