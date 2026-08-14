from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from career.achievement_database import (
    get_achievement_records,
)
from career.database import (
    load_profile,
)
from career.education_database import (
    get_education_records,
)
from career.master_cv import (
    MasterCVData,
    build_master_cv_data,
)
from career.project_database import (
    get_project_records,
)
from career.work_experience_database import (
    get_work_experiences,
)


SECTION_LABELS = {
    "English": {
        "summary": "Professional Summary",
        "skills": "Technical Skills",
        "experience": "Professional Experience",
        "projects": "Projects",
        "education": "Education",
        "achievements": "Selected Achievements",
        "certifications": "Certifications",
        "languages": "Languages",
    },
    "Deutsch": {
        "summary": "Berufliches Profil",
        "skills": "Technische Kenntnisse",
        "experience": "Berufserfahrung",
        "projects": "Projekte",
        "education": "Ausbildung & Studium",
        "achievements": "Ausgewählte Erfolge",
        "certifications": "Zertifikate",
        "languages": "Sprachen",
    },
}


def load_master_cv_data(
    language: str,
) -> MasterCVData:
    """
    Load current verified career records and build the Master CV data.
    """

    return build_master_cv_data(
        profile=load_profile(),
        experiences=get_work_experiences(),
        education_records=(
            get_education_records()
        ),
        projects=get_project_records(),
        achievements=(
            get_achievement_records()
        ),
        language=language,
    )


def _append_text_section(
    parts: list[str],
    heading: str,
    values: list[str],
) -> None:
    if not values:
        return

    rendered = "\n\n".join(
        value.strip()
        for value in values
        if value.strip()
    )

    if rendered:
        parts.append(
            f"{heading}\n{rendered}"
        )


def build_master_cv_text(
    data: MasterCVData,
) -> str:
    labels = SECTION_LABELS[
        data.language
    ]

    parts = []

    if data.candidate_name:
        parts.append(
            data.candidate_name
        )

    if data.professional_title:
        parts.append(
            data.professional_title
        )

    if data.contact_details:
        parts.append(
            " | ".join(
                data.contact_details
            )
        )

    if data.professional_summary:
        parts.append(
            f"{labels['summary']}\n"
            f"{data.professional_summary}"
        )

    if data.technical_skills:
        parts.append(
            f"{labels['skills']}\n"
            + ", ".join(
                data.technical_skills
            )
        )

    _append_text_section(
        parts,
        labels["experience"],
        data.experiences,
    )

    _append_text_section(
        parts,
        labels["projects"],
        data.projects,
    )

    _append_text_section(
        parts,
        labels["education"],
        data.education,
    )

    _append_text_section(
        parts,
        labels["achievements"],
        data.achievements,
    )

    if data.certifications:
        parts.append(
            f"{labels['certifications']}\n"
            + "\n".join(
                f"- {item}"
                for item in data.certifications
            )
        )

    if data.languages:
        parts.append(
            f"{labels['languages']}\n"
            + "\n".join(
                f"- {item}"
                for item in data.languages
            )
        )

    return "\n\n".join(
        part
        for part in parts
        if part.strip()
    )


def _add_multiline_entry(
    document: Document,
    value: str,
) -> None:
    lines = [
        line.strip()
        for line in value.splitlines()
        if line.strip()
    ]

    if not lines:
        return

    first = document.add_paragraph()

    first_run = first.add_run(
        lines[0]
    )

    first_run.bold = True

    for line in lines[1:]:
        if line.startswith("- "):
            document.add_paragraph(
                line[2:],
                style="List Bullet",
            )
        else:
            document.add_paragraph(
                line
            )


def _add_section(
    document: Document,
    heading: str,
    values: list[str],
) -> None:
    if not values:
        return

    document.add_heading(
        heading,
        level=1,
    )

    for value in values:
        if value.strip():
            _add_multiline_entry(
                document,
                value,
            )


def build_master_cv_docx(
    data: MasterCVData,
) -> bytes:
    labels = SECTION_LABELS[
        data.language
    ]

    document = Document()

    normal_style = (
        document.styles["Normal"]
    )

    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(
        10.5
    )

    if data.candidate_name:
        paragraph = (
            document.add_paragraph()
        )

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = paragraph.add_run(
            data.candidate_name
        )

        run.bold = True
        run.font.size = Pt(18)

    if data.professional_title:
        paragraph = (
            document.add_paragraph()
        )

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = paragraph.add_run(
            data.professional_title
        )

        run.bold = True
        run.font.size = Pt(12)

    if data.contact_details:
        paragraph = (
            document.add_paragraph(
                " | ".join(
                    data.contact_details
                )
            )
        )

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

    if data.professional_summary:
        document.add_heading(
            labels["summary"],
            level=1,
        )

        document.add_paragraph(
            data.professional_summary
        )

    if data.technical_skills:
        document.add_heading(
            labels["skills"],
            level=1,
        )

        document.add_paragraph(
            ", ".join(
                data.technical_skills
            )
        )

    _add_section(
        document,
        labels["experience"],
        data.experiences,
    )

    _add_section(
        document,
        labels["projects"],
        data.projects,
    )

    _add_section(
        document,
        labels["education"],
        data.education,
    )

    _add_section(
        document,
        labels["achievements"],
        data.achievements,
    )

    if data.certifications:
        document.add_heading(
            labels["certifications"],
            level=1,
        )

        for item in (
            data.certifications
        ):
            document.add_paragraph(
                item,
                style="List Bullet",
            )

    if data.languages:
        document.add_heading(
            labels["languages"],
            level=1,
        )

        for item in data.languages:
            document.add_paragraph(
                item,
                style="List Bullet",
            )

    output = BytesIO()

    document.save(
        output
    )

    output.seek(0)

    return output.getvalue()


def create_master_cv(
    language: str,
) -> dict:
    data = load_master_cv_data(
        language
    )

    return {
        "data": data,
        "text": build_master_cv_text(
            data
        ),
        "docx": build_master_cv_docx(
            data
        ),
        "warnings": list(
            data.warnings
        ),
    }
