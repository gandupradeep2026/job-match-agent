from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Pt

from career.elevator_pitch import (
    ElevatorPitchRequest,
)
from career.interview_pack import (
    CompleteInterviewPack,
)
from career.interview_preparation import (
    InterviewPreparationRequest,
)
from career.target_company_database import (
    get_target_company,
)
from services.elevator_pitch_service import (
    generate_personal_elevator_pitch,
)
from services.interview_preparation_service import (
    generate_interview_preparation_pack,
)


SECTION_LABELS = {
    "English": {
        "title": "COMPLETE INTERVIEW PACK",
        "pitch": "Personal Elevator Pitch",
        "about": "Tell Me About Yourself",
        "role": "Why This Role?",
        "company": "Why This Company?",
        "strengths": "Strengths",
        "weakness": "Development Area / Weakness",
        "behavioral": "Behavioral / STAR Answers",
        "technical": "Technical Interview Focus",
        "questions": "Questions to Ask the Employer",
        "checklist": "Final Preparation Checklist",
        "warnings": "Warnings",
    },
    "Deutsch": {
        "title": "KOMPLETTES INTERVIEW-PAKET",
        "pitch": "Persönlicher Elevator Pitch",
        "about": "Erzählen Sie mir etwas über sich",
        "role": "Warum diese Position?",
        "company": "Warum dieses Unternehmen?",
        "strengths": "Stärken",
        "weakness": "Entwicklungsbereich / Schwäche",
        "behavioral": "Verhaltensfragen / STAR-Antworten",
        "technical": "Technische Interview-Schwerpunkte",
        "questions": "Fragen an den Arbeitgeber",
        "checklist": "Finale Vorbereitungsliste",
        "warnings": "Hinweise",
    },
}


def _clean(
    values: list[str],
    limit: int | None = None,
) -> list[str]:
    result = []
    seen = set()

    for value in values or []:
        item = str(value).strip()

        if not item:
            continue

        key = item.casefold()

        if key in seen:
            continue

        seen.add(key)
        result.append(item)

        if (
            limit is not None
            and len(result) >= limit
        ):
            break

    return result


def _generate_employer_questions(
    *,
    language: str,
    target_role: str,
    company_id: int | None,
    technical_focus: list[str],
) -> list[str]:
    company = None

    if company_id is not None:
        company = get_target_company(
            company_id
        )

    company_name = (
        company.company_name.strip()
        if company is not None
        else ""
    )

    technologies = []

    if company is not None:
        technologies.extend(
            company.technologies
        )

    technologies.extend(
        technical_focus
    )

    technologies = _clean(
        technologies,
        3,
    )

    if language == "English":
        questions = [
            (
                f"What would success look like in the first six months "
                f"for someone in this {target_role} role?"
                if target_role
                else (
                    "What would success look like in the first six months "
                    "for the person in this role?"
                )
            ),
            (
                "How is the team structured, and how does this role "
                "collaborate with other engineering or business teams?"
            ),
            (
                "What are the most important technical or business "
                "challenges the team is currently working on?"
            ),
            (
                "How do you support learning, feedback and professional "
                "development within the team?"
            ),
            (
                "What are the next steps in the interview process?"
            ),
        ]

        if company_name:
            questions.insert(
                1,
                (
                    f"What are the main priorities for the team at "
                    f"{company_name} over the next 6 to 12 months?"
                ),
            )

        if technologies:
            questions.insert(
                2,
                (
                    "How are "
                    + ", ".join(technologies)
                    + " used in the current technology stack or roadmap?"
                ),
            )

    else:
        questions = [
            (
                f"Wie sieht Erfolg in den ersten sechs Monaten in einer "
                f"Position im Bereich {target_role} aus?"
                if target_role
                else (
                    "Wie sieht Erfolg in den ersten sechs Monaten "
                    "in dieser Position aus?"
                )
            ),
            (
                "Wie ist das Team aufgebaut und wie arbeitet diese Position "
                "mit anderen technischen oder fachlichen Teams zusammen?"
            ),
            (
                "Was sind derzeit die wichtigsten technischen oder "
                "geschäftlichen Herausforderungen des Teams?"
            ),
            (
                "Wie unterstützen Sie Lernen, Feedback und die "
                "fachliche Weiterentwicklung im Team?"
            ),
            (
                "Wie sehen die nächsten Schritte im Bewerbungsprozess aus?"
            ),
        ]

        if company_name:
            questions.insert(
                1,
                (
                    f"Was sind die wichtigsten Prioritäten des Teams bei "
                    f"{company_name} in den kommenden 6 bis 12 Monaten?"
                ),
            )

        if technologies:
            questions.insert(
                2,
                (
                    "Wie werden "
                    + ", ".join(technologies)
                    + " im aktuellen Tech-Stack oder in der technischen "
                    "Roadmap eingesetzt?"
                ),
            )

    return _clean(
        questions,
        7,
    )


def _build_checklist(
    language: str,
) -> list[str]:
    if language == "English":
        return [
            "Review the 30-, 60- and 90-second elevator pitches aloud.",
            "Practice Tell Me About Yourself without reading the text word-for-word.",
            "Review Why This Role and Why This Company before the interview.",
            "Prepare at least three verified STAR stories that cover different competencies.",
            "Review the listed technical focus areas and prepare concrete project examples.",
            "Choose three questions to ask the employer.",
            "Check the company website, job description and interview logistics on the interview day.",
            "Keep all answers truthful and remove anything you cannot confidently support.",
        ]

    return [
        "Die 30-, 60- und 90-Sekunden-Elevator-Pitches laut üben.",
        "Die Selbstvorstellung frei üben, ohne den Text Wort für Wort abzulesen.",
        "Warum diese Position und Warum dieses Unternehmen vor dem Gespräch wiederholen.",
        "Mindestens drei verifizierte STAR-Stories für unterschiedliche Kompetenzen vorbereiten.",
        "Die technischen Schwerpunkte wiederholen und konkrete Projektbeispiele vorbereiten.",
        "Drei Fragen auswählen, die Sie dem Arbeitgeber stellen möchten.",
        "Am Interviewtag Unternehmenswebseite, Stellenbeschreibung und organisatorische Details prüfen.",
        "Alle Antworten wahrheitsgemäß halten und nicht belegbare Aussagen entfernen.",
    ]


def create_complete_interview_pack(
    request: InterviewPreparationRequest,
) -> CompleteInterviewPack:
    if request.language not in (
        "English",
        "Deutsch",
    ):
        raise ValueError(
            "language must be English or Deutsch"
        )

    prep = (
        generate_interview_preparation_pack(
            request
        )
    )

    pitch_results = {}

    for duration in (
        30,
        60,
        90,
    ):
        pitch_results[
            duration
        ] = (
            generate_personal_elevator_pitch(
                ElevatorPitchRequest(
                    language=request.language,
                    duration_seconds=duration,
                    audience="Hiring Manager",
                    target_role=request.target_role,
                )
            )
        )

    warnings = list(
        prep.warnings
    )

    for result in (
        pitch_results.values()
    ):
        warnings.extend(
            result.warnings
        )

    warnings = _clean(
        warnings
    )

    employer_questions = (
        _generate_employer_questions(
            language=request.language,
            target_role=request.target_role,
            company_id=request.company_id,
            technical_focus=(
                prep.technical_focus
            ),
        )
    )

    return CompleteInterviewPack(
        language=request.language,
        target_role=(
            request.target_role
        ),
        company_name=(
            prep.company_name
        ),
        elevator_pitch_30=(
            pitch_results[30].text
        ),
        elevator_pitch_60=(
            pitch_results[60].text
        ),
        elevator_pitch_90=(
            pitch_results[90].text
        ),
        tell_me_about_yourself=(
            prep.tell_me_about_yourself
        ),
        why_this_role=(
            prep.why_this_role
        ),
        why_this_company=(
            prep.why_this_company
        ),
        strengths=(
            prep.strengths
        ),
        weakness=(
            prep.weakness
        ),
        behavioral_answers=(
            prep.behavioral_answers
        ),
        technical_focus=(
            prep.technical_focus
        ),
        employer_questions=(
            employer_questions
        ),
        preparation_checklist=(
            _build_checklist(
                request.language
            )
        ),
        warnings=warnings,
    )


def _answer_text(
    answer,
) -> str:
    if answer is None:
        return ""

    return (
        answer.answer
        or ""
    ).strip()


def build_complete_interview_pack_text(
    pack: CompleteInterviewPack,
) -> str:
    labels = SECTION_LABELS[
        pack.language
    ]

    lines = [
        labels["title"],
        "=" * len(
            labels["title"]
        ),
        "",
        (
            f"Target Role: {pack.target_role}"
            if pack.language == "English"
            else (
                f"Zielrolle: {pack.target_role}"
            )
        ),
    ]

    if pack.company_name:
        lines.append(
            (
                f"Company: {pack.company_name}"
                if pack.language == "English"
                else (
                    f"Unternehmen: {pack.company_name}"
                )
            )
        )

    lines.extend(
        [
            "",
            labels["pitch"].upper(),
            "=" * len(
                labels["pitch"]
            ),
            "",
            "30 seconds / 30 Sekunden",
            pack.elevator_pitch_30,
            "",
            "60 seconds / 60 Sekunden",
            pack.elevator_pitch_60,
            "",
            "90 seconds / 90 Sekunden",
            pack.elevator_pitch_90,
            "",
        ]
    )

    prepared_sections = [
        (
            labels["about"],
            pack.tell_me_about_yourself,
        ),
        (
            labels["role"],
            pack.why_this_role,
        ),
        (
            labels["company"],
            pack.why_this_company,
        ),
        (
            labels["strengths"],
            pack.strengths,
        ),
        (
            labels["weakness"],
            pack.weakness,
        ),
    ]

    for heading, answer in prepared_sections:
        lines.extend(
            [
                heading.upper(),
                "=" * len(
                    heading
                ),
                "",
            ]
        )

        if answer is not None:
            lines.append(
                answer.question
            )

            lines.append(
                _answer_text(
                    answer
                )
                or "[Needs input]"
            )

        lines.append("")

    lines.extend(
        [
            labels["behavioral"].upper(),
            "=" * len(
                labels["behavioral"]
            ),
            "",
        ]
    )

    for index, item in enumerate(
        pack.behavioral_answers,
        start=1,
    ):
        lines.extend(
            [
                f"{index}. {item.category} — {item.story_title}",
                item.question,
                item.answer,
                "",
            ]
        )

    lines.extend(
        [
            labels["technical"].upper(),
            "=" * len(
                labels["technical"]
            ),
        ]
    )

    for item in pack.technical_focus:
        lines.append(
            f"- {item}"
        )

    lines.extend(
        [
            "",
            labels["questions"].upper(),
            "=" * len(
                labels["questions"]
            ),
        ]
    )

    for index, question in enumerate(
        pack.employer_questions,
        start=1,
    ):
        lines.append(
            f"{index}. {question}"
        )

    lines.extend(
        [
            "",
            labels["checklist"].upper(),
            "=" * len(
                labels["checklist"]
            ),
        ]
    )

    for item in (
        pack.preparation_checklist
    ):
        lines.append(
            f"- {item}"
        )

    if pack.warnings:
        lines.extend(
            [
                "",
                labels["warnings"].upper(),
                "=" * len(
                    labels["warnings"]
                ),
            ]
        )

        for warning in pack.warnings:
            lines.append(
                f"- {warning}"
            )

    return "\n".join(
        lines
    )


def _add_heading(
    document: Document,
    text: str,
) -> None:
    document.add_heading(
        text,
        level=1,
    )


def _add_answer(
    document: Document,
    heading: str,
    answer,
) -> None:
    _add_heading(
        document,
        heading,
    )

    if answer is None:
        return

    if answer.question:
        paragraph = (
            document.add_paragraph()
        )

        run = paragraph.add_run(
            answer.question
        )

        run.bold = True

    if answer.answer:
        document.add_paragraph(
            answer.answer
        )


def build_complete_interview_pack_docx(
    pack: CompleteInterviewPack,
) -> bytes:
    labels = SECTION_LABELS[
        pack.language
    ]

    document = Document()

    document.styles[
        "Normal"
    ].font.name = "Arial"

    document.styles[
        "Normal"
    ].font.size = Pt(
        10.5
    )

    title = document.add_heading(
        labels["title"],
        level=0,
    )

    title.alignment = 1

    document.add_paragraph(
        (
            f"Target Role: {pack.target_role}"
            if pack.language == "English"
            else (
                f"Zielrolle: {pack.target_role}"
            )
        )
    )

    if pack.company_name:
        document.add_paragraph(
            (
                f"Company: {pack.company_name}"
                if pack.language == "English"
                else (
                    f"Unternehmen: {pack.company_name}"
                )
            )
        )

    _add_heading(
        document,
        labels["pitch"],
    )

    for duration, text in (
        (
            30,
            pack.elevator_pitch_30,
        ),
        (
            60,
            pack.elevator_pitch_60,
        ),
        (
            90,
            pack.elevator_pitch_90,
        ),
    ):
        paragraph = (
            document.add_paragraph()
        )

        run = paragraph.add_run(
            f"{duration} seconds"
            if pack.language == "English"
            else (
                f"{duration} Sekunden"
            )
        )

        run.bold = True

        document.add_paragraph(
            text
        )

    _add_answer(
        document,
        labels["about"],
        pack.tell_me_about_yourself,
    )

    _add_answer(
        document,
        labels["role"],
        pack.why_this_role,
    )

    _add_answer(
        document,
        labels["company"],
        pack.why_this_company,
    )

    _add_answer(
        document,
        labels["strengths"],
        pack.strengths,
    )

    _add_answer(
        document,
        labels["weakness"],
        pack.weakness,
    )

    _add_heading(
        document,
        labels["behavioral"],
    )

    for item in (
        pack.behavioral_answers
    ):
        paragraph = (
            document.add_paragraph()
        )

        run = paragraph.add_run(
            (
                f"{item.category} — "
                f"{item.story_title}"
            )
        )

        run.bold = True

        document.add_paragraph(
            item.question
        )

        for line in (
            item.answer.splitlines()
        ):
            if line.strip():
                document.add_paragraph(
                    line.strip()
                )

    _add_heading(
        document,
        labels["technical"],
    )

    for item in pack.technical_focus:
        document.add_paragraph(
            item,
            style="List Bullet",
        )

    _add_heading(
        document,
        labels["questions"],
    )

    for item in (
        pack.employer_questions
    ):
        document.add_paragraph(
            item,
            style="List Number",
        )

    _add_heading(
        document,
        labels["checklist"],
    )

    for item in (
        pack.preparation_checklist
    ):
        document.add_paragraph(
            item,
            style="List Bullet",
        )

    if pack.warnings:
        _add_heading(
            document,
            labels["warnings"],
        )

        for item in pack.warnings:
            document.add_paragraph(
                item,
                style="List Bullet",
            )

    output = BytesIO()

    document.save(
        output
    )

    output.seek(0)

    return output.getvalue()


def generate_complete_interview_pack_files(
    request: InterviewPreparationRequest,
) -> dict:
    pack = (
        create_complete_interview_pack(
            request
        )
    )

    return {
        "pack": pack,
        "text": (
            build_complete_interview_pack_text(
                pack
            )
        ),
        "docx": (
            build_complete_interview_pack_docx(
                pack
            )
        ),
    }
