from io import BytesIO
from pathlib import Path
from typing import Any

import fitz
from docx import Document

from services.ocr_service import (
    OCRError,
    extract_text_from_pdf_with_ocr,
)


MIN_DIRECT_PDF_TEXT_LENGTH = 80


def read_uploaded_file_bytes(
    uploaded_file,
) -> bytes:
    """
    Read uploaded file bytes safely.
    """

    if hasattr(
        uploaded_file,
        "getvalue",
    ):
        return uploaded_file.getvalue()

    if hasattr(
        uploaded_file,
        "read",
    ):
        file_bytes = uploaded_file.read()

        try:
            uploaded_file.seek(0)

        except Exception:
            pass

        return file_bytes

    raise ValueError(
        "The uploaded file could not be read."
    )


def extract_text_from_pdf_bytes(
    pdf_bytes: bytes,
) -> str:
    """
    Extract selectable text from a PDF.
    """

    if not pdf_bytes:
        return ""

    try:
        document = fitz.open(
            stream=pdf_bytes,
            filetype="pdf",
        )

    except Exception as error:
        raise ValueError(
            "The uploaded PDF could not be opened."
        ) from error

    try:
        page_texts = []

        for page in document:
            page_texts.append(
                page.get_text(
                    "text"
                )
            )

        return "\n".join(
            page_texts
        ).strip()

    finally:
        document.close()


def extract_text_from_docx_bytes(
    docx_bytes: bytes,
) -> str:
    """
    Extract paragraph and table text from a DOCX file.
    """

    if not docx_bytes:
        return ""

    try:
        document = Document(
            BytesIO(
                docx_bytes
            )
        )

    except Exception as error:
        raise ValueError(
            "The uploaded DOCX file could not be opened."
        ) from error

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    table_text = []

    for table in document.tables:
        for row in table.rows:
            values = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]

            if values:
                table_text.append(
                    " | ".join(
                        values
                    )
                )

    return "\n".join(
        paragraphs
        + table_text
    ).strip()


def extract_text_from_txt_bytes(
    text_bytes: bytes,
) -> str:
    """
    Decode a TXT file using common encodings.
    """

    if not text_bytes:
        return ""

    for encoding in [
        "utf-8",
        "utf-8-sig",
        "cp1252",
        "latin-1",
    ]:
        try:
            return text_bytes.decode(
                encoding
            ).strip()

        except UnicodeDecodeError:
            continue

    return text_bytes.decode(
        "utf-8",
        errors="replace",
    ).strip()


def extract_pdf_text_with_fallback(
    pdf_bytes: bytes,
) -> tuple[str, dict[str, Any]]:
    """
    Extract PDF text directly and use OCR when needed.
    """

    direct_text = extract_text_from_pdf_bytes(
        pdf_bytes
    )

    if (
        len(
            direct_text.strip()
        )
        >= MIN_DIRECT_PDF_TEXT_LENGTH
    ):
        return (
            direct_text,
            {
                "method": "direct_pdf_text",
                "ocr_used": False,
                "warnings": [],
                "page_count": 0,
                "processed_pages": 0,
                "failed_pages": [],
                "languages": "",
                "dpi": 0,
            },
        )

    try:
        ocr_result = (
            extract_text_from_pdf_with_ocr(
                pdf_bytes=pdf_bytes,
                languages="eng+deu",
            )
        )

        return (
            ocr_result["text"],
            {
                "method": "ocr",
                "ocr_used": True,
                "warnings": (
                    ocr_result.get(
                        "warnings",
                        [],
                    )
                ),
                "page_count": (
                    ocr_result.get(
                        "page_count",
                        0,
                    )
                ),
                "processed_pages": (
                    ocr_result.get(
                        "processed_pages",
                        0,
                    )
                ),
                "failed_pages": (
                    ocr_result.get(
                        "failed_pages",
                        [],
                    )
                ),
                "languages": (
                    ocr_result.get(
                        "languages",
                        "eng+deu",
                    )
                ),
                "dpi": (
                    ocr_result.get(
                        "dpi",
                        0,
                    )
                ),
            },
        )

    except OCRError as error:
        if direct_text.strip():
            return (
                direct_text,
                {
                    "method": "direct_pdf_text",
                    "ocr_used": False,
                    "warnings": [
                        (
                            "OCR fallback failed: "
                            f"{error}"
                        )
                    ],
                    "page_count": 0,
                    "processed_pages": 0,
                    "failed_pages": [],
                    "languages": "",
                    "dpi": 0,
                },
            )

        raise ValueError(
            "No readable text could be extracted "
            "from the PDF, and OCR also failed: "
            f"{error}"
        ) from error


def extract_document_text_with_details(
    uploaded_file,
) -> dict[str, Any]:
    """
    Extract document text and return extraction diagnostics.
    """

    if uploaded_file is None:
        raise ValueError(
            "No file was provided."
        )

    filename = getattr(
        uploaded_file,
        "name",
        "",
    )

    file_extension = Path(
        filename
    ).suffix.lower()

    file_bytes = read_uploaded_file_bytes(
        uploaded_file
    )

    if not file_bytes:
        raise ValueError(
            "The uploaded file is empty."
        )

    if file_extension == ".pdf":
        extracted_text, details = (
            extract_pdf_text_with_fallback(
                file_bytes
            )
        )

        return {
            "text": extracted_text,
            "details": {
                **details,
                "file_type": "pdf",
                "filename": filename,
            },
        }

    if file_extension == ".docx":
        extracted_text = (
            extract_text_from_docx_bytes(
                file_bytes
            )
        )

        return {
            "text": extracted_text,
            "details": {
                "method": "docx",
                "ocr_used": False,
                "warnings": [],
                "file_type": "docx",
                "filename": filename,
                "page_count": 0,
                "processed_pages": 0,
                "failed_pages": [],
                "languages": "",
                "dpi": 0,
            },
        }

    if file_extension == ".txt":
        extracted_text = (
            extract_text_from_txt_bytes(
                file_bytes
            )
        )

        return {
            "text": extracted_text,
            "details": {
                "method": "txt",
                "ocr_used": False,
                "warnings": [],
                "file_type": "txt",
                "filename": filename,
                "page_count": 0,
                "processed_pages": 0,
                "failed_pages": [],
                "languages": "",
                "dpi": 0,
            },
        }

    raise ValueError(
        "Unsupported file type. "
        "Upload a PDF, DOCX or TXT file."
    )


def extract_document_text(
    uploaded_file,
) -> str:
    """
    Backward-compatible text-only extraction.
    """

    result = (
        extract_document_text_with_details(
            uploaded_file
        )
    )

    return result["text"]